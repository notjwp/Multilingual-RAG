from pathlib import Path

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from multilingual_rag.core.config import Settings
from multilingual_rag.core.errors import AppError
from multilingual_rag.ingestion.language import LanguageDetector
from multilingual_rag.ingestion.loaders import DocumentLoader
from multilingual_rag.ingestion.service import IngestionService


def test_language_detector_returns_default_for_short_text() -> None:
    detector = LanguageDetector(min_text_length=20)

    assert detector.detect("short", default="unknown") == "unknown"


def test_language_detector_detects_english() -> None:
    detector = LanguageDetector(min_text_length=0)

    assert detector.detect("This is a clear English sentence about retrieval systems.") == "en"


def test_text_ingestion_produces_document_metadata_and_chunks(tmp_path: Path) -> None:
    source_path = tmp_path / "english.txt"
    source_path.write_text(
        "This is a multilingual retrieval augmented generation document. " * 4,
        encoding="utf-8",
    )
    service = IngestionService(Settings(chunk_size_tokens=12, chunk_overlap_tokens=2))

    result = service.ingest_file(source_path)

    assert result.document.source == str(source_path.resolve())
    assert result.document.content_type == "text/plain"
    assert result.document.language == "en"
    assert len(result.chunks) >= 2
    assert all(chunk.document_id == result.document.document_id for chunk in result.chunks)


def test_html_loader_removes_script_content(tmp_path: Path) -> None:
    source_path = tmp_path / "page.html"
    source_path.write_text(
        "<html><script>hidden()</script><body><h1>Hello</h1><p>Visible text</p></body></html>",
        encoding="utf-8",
    )

    loaded_document = DocumentLoader().load(source_path)

    assert "Visible text" in loaded_document.sections[0].text
    assert "hidden" not in loaded_document.sections[0].text


def test_docx_loader_extracts_paragraphs(tmp_path: Path) -> None:
    source_path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.add_paragraph("First paragraph")
    document.add_paragraph("Second paragraph")
    document.save(source_path)

    loaded_document = DocumentLoader().load(source_path)

    assert [section.text for section in loaded_document.sections] == [
        "First paragraph",
        "Second paragraph",
    ]


def test_pdf_loader_extracts_page_sections(tmp_path: Path) -> None:
    source_path = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with source_path.open("wb") as file:
        writer.write(file)

    loaded_document = DocumentLoader().load(source_path)

    assert loaded_document.content_type == "application/pdf"
    assert loaded_document.metadata["page_count"] == 1
    assert loaded_document.sections[0].page == 1


def test_loader_rejects_unsupported_file_type(tmp_path: Path) -> None:
    source_path = tmp_path / "data.csv"
    source_path.write_text("a,b,c", encoding="utf-8")

    with pytest.raises(AppError, match="Unsupported document type"):
        DocumentLoader().load(source_path)


def test_ingestion_rejects_empty_document(tmp_path: Path) -> None:
    source_path = tmp_path / "empty.txt"
    source_path.write_text("   \n\n", encoding="utf-8")
    service = IngestionService(Settings())

    with pytest.raises(AppError, match="extractable text"):
        service.ingest_file(source_path)
