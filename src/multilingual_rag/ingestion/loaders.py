"""Document loaders for supported source formats."""

from __future__ import annotations

from collections.abc import Callable
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from multilingual_rag.core.errors import AppError
from multilingual_rag.core.models import DocumentSection, LoadedDocument

Loader = Callable[[Path], LoadedDocument]


CONTENT_TYPES: dict[str, str] = {
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".markdown": "text/markdown",
    ".html": "text/html",
    ".htm": "text/html",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


class DocumentLoader:
    """Load supported document files into text sections."""

    def load(self, path: Path) -> LoadedDocument:
        """Parse a document file based on its extension."""
        resolved_path = path.expanduser().resolve()
        if not resolved_path.exists():
            raise AppError(
                f"Document does not exist: {path}",
                code="document_not_found",
                status_code=404,
            )
        if not resolved_path.is_file():
            raise AppError(
                f"Document path is not a file: {path}",
                code="invalid_document_path",
                status_code=400,
            )

        suffix = resolved_path.suffix.lower()
        loader = self._loader_for_suffix(suffix)
        return loader(resolved_path)

    def _loader_for_suffix(self, suffix: str) -> Loader:
        loaders: dict[str, Loader] = {
            ".txt": load_text_document,
            ".md": load_text_document,
            ".markdown": load_text_document,
            ".html": load_html_document,
            ".htm": load_html_document,
            ".pdf": load_pdf_document,
            ".docx": load_docx_document,
        }
        try:
            return loaders[suffix]
        except KeyError as exc:
            raise AppError(
                f"Unsupported document type: {suffix or '<none>'}",
                code="unsupported_document_type",
                status_code=400,
            ) from exc


def load_text_document(path: Path) -> LoadedDocument:
    """Load plain text or Markdown documents."""
    text = read_text_with_fallback(path)
    section = DocumentSection(text=text, section_index=0)
    return LoadedDocument(
        source_path=path,
        content_type=CONTENT_TYPES[path.suffix.lower()],
        sections=(section,),
    )


def load_html_document(path: Path) -> LoadedDocument:
    """Load HTML documents as visible text."""
    html = read_text_with_fallback(path)
    soup = BeautifulSoup(html, "html.parser")

    for element in soup(["script", "style", "noscript"]):
        element.decompose()

    text = soup.get_text(separator="\n")
    section = DocumentSection(text=text, section_index=0)
    return LoadedDocument(
        source_path=path,
        content_type=CONTENT_TYPES[path.suffix.lower()],
        sections=(section,),
    )


def load_pdf_document(path: Path) -> LoadedDocument:
    """Load text from each page of a PDF document."""
    reader = PdfReader(path)
    sections: list[DocumentSection] = []

    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        sections.append(
            DocumentSection(
                text=text,
                page=index,
                section_index=index - 1,
            )
        )

    return LoadedDocument(
        source_path=path,
        content_type=CONTENT_TYPES[path.suffix.lower()],
        sections=tuple(sections),
        metadata={"page_count": len(sections)},
    )


def load_docx_document(path: Path) -> LoadedDocument:
    """Load text from DOCX paragraphs."""
    document = DocxDocument(str(path))
    sections = tuple(
        DocumentSection(text=paragraph.text, section_index=index)
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip()
    )
    return LoadedDocument(
        source_path=path,
        content_type=CONTENT_TYPES[path.suffix.lower()],
        sections=sections,
        metadata={"paragraph_count": len(sections)},
    )


def read_text_with_fallback(path: Path) -> str:
    """Read text with a small set of common encodings."""
    encodings = ("utf-8", "utf-8-sig", "cp1252")
    last_error: UnicodeDecodeError | None = None

    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as exc:
            last_error = exc

    raise AppError(
        f"Unable to decode text document: {path}",
        code="document_decode_error",
        status_code=400,
    ) from last_error
